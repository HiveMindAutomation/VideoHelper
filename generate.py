#!/bin/python3

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

import json
import os
import requests

## This is a local file containing some "Sensitive" data, but also the URL's Below
import auth

# Using JIRA integration?
# Change to False if not
JIRA = True

rootPath = "/Volumes/HiveMind/Videos"

#### Variables ####
## I've imported a seperate file called auth.py to keep this a bit more secure
## Replace the F-Strings in here with your own Details, or create your own auth.py 
facebookURL = f'{auth.facebookURL}'
twitterURL = f'{auth.twitterURL}'
instagramURL = f'{auth.instagramURL}'

name = f'{auth.name}'
channelName = f'{auth.channelName}'
channelAddress = f'{auth.channelAddress}'
project = f'{auth.project}'

############ JIRA ###########

## JIRA URL
if JIRA == True:
    # JIRA URL and Headers pulled from Auth.py file
    url = f'{auth.url}'\
    # Authentication is needed in your request Header to access the JIRA API
    headers = auth.headers
else:
    pass

## This defines the prompt we see when selecting the "TYPE" of project
## Triple ' marks keeps the string open across multiple lines. This (theoretically) makes the longer strings below a lot easier to edit
typePrompt = '''What type of Project is this?
1 - Getting Started
2 - Product Review
3 - Quick Tips
4 - Code Review

?:  '''

# Define video Type
# Asks the user to select Video Type based on the above
projectType = 0
while projectType not in [1,2,3,4]:
    ## IF Statement to determine where to save the files based on the Project Type selected above
    projectType = int(input(typePrompt))

# Ask user to Enter the JIRA Project ID
projectID = input(f"What is the Project ID after {project}- ? : ")

if JIRA == True:
    # Pull Video Title from JIRA Project using JIRA API
    result = requests.request("GET", f"{url}/rest/api/latest/issue/{project}-{projectID}", verify=False, headers=headers)
    # Parse the result into JSON which we can reference like dicts.
    JSONResult = json.loads(result.text)

    #Output Video Title to Terminal - Mostly for Error checking but kinda nice to have
    print(JSONResult['fields']['summary'])
    # Assign JIRA Project Summary to Job Title
    jobTitle = JSONResult['fields']['summary']
else:
    jobTitle = str(input("What is the title of this project?: "))

# Set Path based on inputs
if projectType == 1:
    path = f"{rootPath}/Getting Started Series/{project}-{projectID} - {jobTitle}"
elif projectType == 2:
    path = f"{rootPath}/Product Reviews/{project}-{projectID} - {jobTitle}"
elif projectType == 3:
    path = f"{rootPath}/Quick Tips/{project}-{projectID} - {jobTitle}"
elif projectType == 4:
    path = f"{rootPath}/Code Review/{project}-{projectID} - {jobTitle}"
else:
    projectType = int(input(typePrompt))

## I render my videos out into a path called "Render" in a subfolder under each project.
## All files associated with the YouTube upload go into this folder
renderpath = f"{path}/Render"

## Create Project Directories
## You could reasonably create any Directory structure you want here.
os.mkdir(path)
## Create Render Directory
os.mkdir(renderpath)

# TODO - Mount SMB Share


################    TEMPLATES   ################ 

## Templated Intro
introTemplate = f"""Hi, I'm {name} from {channelName} and welcome to the Hive!


In This Video we'll be taking a look at {jobTitle}.
.
.
.
.
.
.

While I roll the intro, take a moment to Subscribe, and hit the bell icon to get notified when I release new videos each week.

Let's get started!"""

## Templated Outtro
outtroTemplate = f"""That's all we have for this video and I hope it helped you in your home automation journey.

Be sure to comment down below with a home automation idea you'd like to see me cover in a future video.
Don't forget to Follow {channelName} on Twitter, Instagram and Facebook.

If you liked this video, hit the Thumbs Up button down below to give it a like.

And if you're not already subscribed, please consider subscribing now.
While you're at it, hit the bell icon to get notified when I release new videos each week.

Lastly, if you like what I'm doing here, and you want to help support the channel, there's a buy me a coffee link in the video description below.

Contributions through Buy me a coffee are put towards making more, and better content for you to enjoy.

Thanks so much for watching! I'm {name} from {channelName}
And I'm looking forward to seeing you next time!

Bye for now!"""

## Template for YouTube Description
descriptionTextBase = f"""{jobTitle}

*** Links ***

{channelName} on YouTube: {channelAddress}

*** Support the Channel***
Buy Me a Coffee: https://buymeacoffee.com/HiveMindAuto

*** Find Hive Mind Automation on Social Media ***

Twitter: {twitterURL}
Instagram: {instagramURL}
Facebook: {facebookURL}

*** Affiliate Links ***
*** These links help the channel by providing a commission on purchases\n\n
*** TIMESTAMPS ***

0:00 Intro


*** Helpful Links ***

Home Assistant: https://www.home-assistant.io/
Raspberry Pi: https://www.raspberrypi.org/
Balena Etcher: https://www.balena.io/etcher/

Home Assistant for iOS: https://apple.co/34JATce
Home Assistant for Android: https://bit.ly/30VUsNh

*** CREDITS ***

Music: https://www.purple-planet.com
"""

###########################     Create the WORD Document    ############################
# Instantiate Document
document = Document()

# Insert Header
heading = document.add_heading(f"{project}-{projectID}:\n{jobTitle}", 0)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Insert Intro Text Block
paragraph = document.add_paragraph("<INTRO>")
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
## This is Where the Template gets inserted
paragraph = document.add_paragraph(introTemplate)
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph = document.add_paragraph("<ROLL INTRO ANIMATION>")
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
## All this other stuff is just getting the document laid out in a useful way
paragraph = document.add_paragraph("\n\n<PREFACE>")
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph = document.add_paragraph("\n.\n.\n")
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Insert Summary Text Block
paragraph = document.add_paragraph("<SUMMARY>")
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph = document.add_paragraph(".\n.\n.\n.\n.")
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Insert Outtro Text Block 
paragraph = document.add_paragraph("<OUTTRO>")
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
## Actual Outtro Template inserts here
paragraph = document.add_paragraph(outtroTemplate)
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph = document.add_paragraph()
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph = document.add_paragraph("<CUT>")
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save .docx File in project path
document.save(f"{path}/{project}-{projectID} - {jobTitle}.docx")
#############################################################################################


######## Create Text File with "Description" Template
# Generate YouTube Description Template
descriptionFile = open(f"{renderpath}/{project}-{projectID}-{jobTitle} - YouTube Description.txt", "w")
## Write Description Text out to file
descriptionFile.write(descriptionTextBase)
## Close the description File
descriptionFile.close()

###################################     END     ################################### 